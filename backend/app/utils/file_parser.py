"""Resume and document parsing utilities with comprehensive logging.

Supports PDF and DOCX files, extracting raw text for downstream processing.
Logs: file type, page counts, paragraph counts, text sizes, and timing.
"""

import logging
import time
from pathlib import Path

logger = logging.getLogger("hireai.file_parser")


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    start = time.time()
    logger.info(f"FILE_PARSER.pdf | START | path={file_path}")

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        page_count = len(doc)
        logger.info(f"FILE_PARSER.pdf | Document opened | pages={page_count}")

        text_parts = []
        for i, page in enumerate(doc):
            page_text = page.get_text()
            text_parts.append(page_text)
            logger.debug(
                f"FILE_PARSER.pdf | Page {i+1}/{page_count}"
                f" | chars={len(page_text)}"
            )
        doc.close()

        full_text = "\n".join(text_parts).strip()
        elapsed = int((time.time() - start) * 1000)
        logger.info(
            f"FILE_PARSER.pdf | COMPLETE"
            f" | pages={page_count}"
            f" | total_chars={len(full_text)}"
            f" | {elapsed}ms"
        )
        return full_text

    except Exception as e:
        elapsed = int((time.time() - start) * 1000)
        logger.error(
            f"FILE_PARSER.pdf | FAILED"
            f" | path={file_path}"
            f" | {elapsed}ms"
            f" | {type(e).__name__}: {e}"
        )
        raise ValueError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    start = time.time()
    logger.info(f"FILE_PARSER.docx | START | path={file_path}")

    try:
        from docx import Document
        doc = Document(file_path)
        total_paragraphs = len(doc.paragraphs)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        full_text = "\n".join(text_parts).strip()
        elapsed = int((time.time() - start) * 1000)
        logger.info(
            f"FILE_PARSER.docx | COMPLETE"
            f" | total_paragraphs={total_paragraphs}"
            f" | non_empty_paragraphs={len(text_parts)}"
            f" | total_chars={len(full_text)}"
            f" | {elapsed}ms"
        )
        return full_text

    except Exception as e:
        elapsed = int((time.time() - start) * 1000)
        logger.error(
            f"FILE_PARSER.docx | FAILED"
            f" | path={file_path}"
            f" | {elapsed}ms"
            f" | {type(e).__name__}: {e}"
        )
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_file(file_path: str) -> str:
    """Extract text from a file based on its extension."""
    path = Path(file_path)
    ext = path.suffix.lower()
    file_size_kb = path.stat().st_size / 1024

    logger.info(
        f"FILE_PARSER.dispatch | Extracting text"
        f" | path={path.name}"
        f" | type={ext}"
        f" | size={file_size_kb:.1f}KB"
    )

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif ext == ".txt":
        text = path.read_text(encoding="utf-8")
        logger.info(
            f"FILE_PARSER.txt | COMPLETE"
            f" | chars={len(text)}"
        )
        return text
    else:
        logger.error(f"FILE_PARSER.dispatch | Unsupported file type: {ext}")
        raise ValueError(f"Unsupported file type: {ext}. Supported: .pdf, .docx, .txt")
